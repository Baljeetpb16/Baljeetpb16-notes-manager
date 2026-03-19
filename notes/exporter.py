"""
Note export utilities.

Supports exporting a note's metadata and text content to:
  - Markdown (.md)
  - PDF     (.pdf)  – via fpdf2
  - DOCX    (.docx) – via python-docx
"""

from __future__ import annotations

import io
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from notes.models import Note


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------


def export_markdown(note: "Note") -> bytes:
    """Return a Markdown representation of *note* as UTF-8 bytes."""
    lines = [
        f"# {note.title}",
        "",
        f"**Subject:** {note.subject}",
        f"**Tags:** {note.tags or '—'}",
        f"**Visibility:** {note.get_visibility_display()}",
        f"**Created:** {note.created_at.strftime('%Y-%m-%d %H:%M')}",
        "",
    ]
    if note.content:
        lines += ["## Content", "", note.content, ""]
    return "\n".join(lines).encode("utf-8")


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def export_pdf(note: "Note") -> bytes:
    """Return a PDF representation of *note* as bytes."""
    from fpdf import FPDF  # type: ignore[import-untyped]

    def _safe(text: str) -> str:
        """Replace characters unsupported by the latin-1 core fonts with safe equivalents."""
        return (
            text.replace("\u2014", "--")  # em dash
            .replace("\u2013", "-")  # en dash
            .replace("\u2018", "'")  # left single quote
            .replace("\u2019", "'")  # right single quote
            .replace("\u201c", '"')  # left double quote
            .replace("\u201d", '"')  # right double quote
            .encode("latin-1", errors="replace")
            .decode("latin-1")
        )

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Title
    pdf.set_font("Helvetica", style="B", size=18)
    pdf.cell(0, 12, _safe(note.title), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # Metadata
    tags_display = note.tags if note.tags else "-"
    meta = [
        ("Subject", _safe(str(note.subject))),
        ("Tags", _safe(tags_display)),
        ("Visibility", _safe(note.get_visibility_display())),
        ("Created", note.created_at.strftime("%Y-%m-%d %H:%M")),
    ]
    for label, value in meta:
        pdf.set_font("Helvetica", style="B", size=11)
        pdf.cell(30, 8, f"{label}:", new_x="END", new_y="LAST")
        pdf.set_font("Helvetica", size=11)
        pdf.cell(0, 8, value, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Content
    if note.content:
        pdf.set_font("Helvetica", style="B", size=13)
        pdf.cell(0, 10, "Content", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", size=11)
        pdf.multi_cell(0, 7, _safe(note.content))

    return bytes(pdf.output())


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def export_docx(note: "Note") -> bytes:
    """Return a DOCX representation of *note* as bytes."""
    from docx import Document  # type: ignore[import-untyped]
    from docx.shared import Pt  # type: ignore[import-untyped]

    doc = Document()

    doc.add_heading(note.title, level=1)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("Subject", str(note.subject)),
        ("Tags", note.tags or "—"),
        ("Visibility", note.get_visibility_display()),
        ("Created", note.created_at.strftime("%Y-%m-%d %H:%M")),
    ]
    for row, (label, value) in zip(table.rows, rows_data):
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)

    doc.add_paragraph()

    if note.content:
        doc.add_heading("Content", level=2)
        doc.add_paragraph(note.content)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
